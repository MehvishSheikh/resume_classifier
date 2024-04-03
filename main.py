# import streamlit as st
# import pickle
# import re
# from PyPDF2 import PdfReader
# from docx import Document

# # Load the TF-IDF vectorizer
# with open('tfidf.pkl', 'rb') as f:
#     loaded_tfidf = pickle.load(f)

# # Load the classifier
# with open('clf.pkl', 'rb') as f:
#     loaded_clf = pickle.load(f)

# # Define function to clean resume text
# def clean_text(text):
#     # Remove extra spaces
#     cleaned_text = re.sub(r'\s+', ' ', text)
#     return cleaned_text

# # Define the main function to classify the resume
# def classify_resume(uploaded_file):
#     # Initialize text variable
#     text = ""

#     # Check file type and read content
#     if uploaded_file.type == 'application/pdf':
#         # Read PDF file
#         pdf_reader = PdfReader(uploaded_file)
#         for page in pdf_reader.pages:
#             text += page.extract_text()
#     elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
#         # Read DOCX file
#         doc = Document(uploaded_file)
#         for para in doc.paragraphs:
#             text += para.text + '\n'

#     # Clean the text
#     cleaned_text = clean_text(text)

#     # Transform the cleaned text using the loaded TF-IDF vectorizer
#     resume_tfidf = loaded_tfidf.transform([cleaned_text])

#     # Predict the category of the resume using the loaded classifier
#     predicted_category = loaded_clf.predict(resume_tfidf)[0]

#     return predicted_category

# # Streamlit UI
# def main():
#     st.title("Resume Classifier")

#     # File uploader widget
#     uploaded_file = st.file_uploader("Upload a resume (PDF or Word)", type=['pdf', 'docx'])

#     if uploaded_file is not None:
#         # Display the uploaded file
#         st.write("Uploaded resume:")
#         st.write(uploaded_file)

#         # Classify the resume
#         classification_result = classify_resume(uploaded_file)

#         # Display the classification result
#         st.write("Classification result:")
#         st.write("Predicted Category:", classification_result)

# if __name__ == "__main__":
#     main()
import streamlit as st
import pickle
import re
from PyPDF2 import PdfReader
from docx import Document

# Load the TF-IDF vectorizer
with open('tfidf.pkl', 'rb') as f:
    loaded_tfidf = pickle.load(f)

# Load the classifier
with open('clf.pkl', 'rb') as f:
    loaded_clf = pickle.load(f)

# Define function to clean resume text
def clean_text(text):
    # Remove extra spaces
    cleaned_text = re.sub(r'\s+', ' ', text)
    return cleaned_text

# List of class names
classes = ['Data Science', 'HR', 'Advocate', 'Arts', 'Web Designing',
           'Mechanical Engineer', 'Sales', 'Health and fitness',
           'Civil Engineer', 'Java Developer', 'Business Analyst',
           'SAP Developer', 'Automation Testing', 'Electrical Engineering',
           'Operations Manager', 'Python Developer', 'DevOps Engineer',
           'Network Security Engineer', 'PMO', 'Database', 'Hadoop',
           'ETL Developer', 'DotNet Developer', 'Blockchain', 'Testing']

# Example mapping of category integers to class names
category_mapping = {
    0: 'Data Science',
    1: 'HR',
    2: 'Advocate',
    3: 'Arts',
    4: 'Web Designing',
    5: 'Mechanical Engineer',
    6: 'Sales',
    7: 'Health and fitness',
    8: 'Civil Engineer',
    9: 'Java Developer',
    10: 'Business Analyst',
    11: 'SAP Developer',
    12: 'Automation Testing',
    13: 'Electrical Engineering',
    14: 'Operations Manager',
    15: 'Python Developer',
    16: 'DevOps Engineer',
    17: 'Network Security Engineer',
    18: 'PMO',
    19: 'Database',
    20: 'Hadoop',
    21: 'ETL Developer',
    22: 'DotNet Developer',
    23: 'Blockchain',
    24: 'Testing'
}

# Define the main function to classify the resume
def classify_resume(uploaded_file):
    # Initialize text variable
    text = ""

    # Check file type and read content
    if uploaded_file.type == 'application/pdf':
        # Read PDF file
        pdf_reader = PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        # Read DOCX file
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + '\n'

    # Clean the text
    cleaned_text = clean_text(text)

    # Transform the cleaned text using the loaded TF-IDF vectorizer
    resume_tfidf = loaded_tfidf.transform([cleaned_text])

    # Predict the category of the resume using the loaded classifier
    predicted_category = loaded_clf.predict(resume_tfidf)[0]

    # Get the class name from the category integer
    class_name = category_mapping[predicted_category]

    return predicted_category, class_name

# Streamlit UI
def main():
    st.title("Resume Classifier")

    # File uploader widget
    uploaded_file = st.file_uploader("Upload a resume (PDF or Word)", type=['pdf', 'docx'])

    if uploaded_file is not None:
        # Display the uploaded file
        st.write("Uploaded resume:")
        st.write(uploaded_file)

        # Classify the resume
        category_int, class_name = classify_resume(uploaded_file)

        # Display the classification result
        st.write("Classification result:")
        st.write("Predicted Category:", category_int)
        st.write("Predicted Class:", class_name)

if __name__ == "__main__":
    main()
